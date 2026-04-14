from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os
import logging

log = logging.getLogger(__name__)
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'reports')
os.makedirs(OUTPUT_DIR, exist_ok=True)


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
        r.font.size = Pt(13)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x2E, 0x6D, 0xA4)
    return p


def kv_field(doc, label, value):
    if not value:
        return
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
    p.add_run(str(value))


def add_bullet(doc, text):
    doc.add_paragraph(style='List Bullet').add_run(str(text))


def generate_word_report(data: dict) -> str:
    esc = data.get('escalation_summary', {})
    inv = data.get('investigation', {})
    rca = data.get('rca', {})
    vt = data.get('vt_enrichment', {})
    ts = data.get('timestamp', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1.2)
    sec.right_margin = Inches(1.2)

    # Title
    t = doc.add_heading('ROOT CAUSE ANALYSIS REPORT', 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in t.runs:
        r.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
        r.font.size = Pt(20)

    sub = doc.add_paragraph('ThreatMail AI — CTI Escalation Investigation Platform')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    gen = doc.add_paragraph(f'Generated: {ts}   |   Confidential — For Authorized Use Only')
    gen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    gen.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # 1. Problem Statement
    h1(doc, '1. Problem Statement')
    tbl = doc.add_table(rows=7, cols=2)
    tbl.style = 'Table Grid'
    # Header row
    hr = tbl.rows[0]
    for i, label in enumerate(['Field', 'Details']):
        hr.cells[i].text = label
        hr.cells[i].paragraphs[0].runs[0].bold = True
        hr.cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hr.cells[i], '1A3A5C')
    rows_data = [
        ('Problem Statement', rca.get('problem_statement', 'N/A')),
        ('Report Month', ts[:7]),
        ('Escalation Type', (esc.get('escalation_type', '') or '').replace('_', ' ').title()),
        ('Severity', esc.get('severity', 'N/A')),
        ('Platform Affected', esc.get('platform_affected', 'N/A')),
        ('Brand / Client', esc.get('brand_targeted', 'N/A')),
    ]
    for i, (label, value) in enumerate(rows_data):
        r = tbl.rows[i + 1]
        r.cells[0].text = label
        r.cells[0].paragraphs[0].runs[0].bold = True
        r.cells[1].text = str(value)
    doc.add_paragraph()

    # 2. Executive Summary
    h1(doc, '2. Executive Summary')
    doc.add_paragraph(rca.get('executive_summary', 'N/A'))
    doc.add_paragraph()

    # 3. Escalation Details
    h1(doc, '3. Escalation Details')
    kv_field(doc, 'Escalation Summary', esc.get('escalation_summary'))
    kv_field(doc, 'Pillar Affected', esc.get('pillar_affected'))
    kv_field(doc, 'Detection Issue', esc.get('detection_issue'))
    kv_field(doc, 'Client Impact', esc.get('client_impact'))
    kv_field(doc, 'Ticket Reference', esc.get('ticket_reference'))
    doc.add_paragraph()

    # 4. Assets
    h1(doc, '4. Assets Identified')
    from modules.virustotal import extract_domain
    for asset in inv.get('assets_investigated', []):
        h2(doc, asset.get('asset', 'Unknown Asset'))
        kv_field(doc, 'Type', asset.get('asset_type'))
        kv_field(doc, 'Risk Level', asset.get('risk_level'))
        kv_field(doc, 'Likely Purpose', asset.get('likely_purpose'))
        kv_field(doc, 'Why Suspicious', asset.get('why_suspicious'))
        for ind in (asset.get('indicators') or []):
            add_bullet(doc, ind)
        # VT data
        dk = extract_domain(asset.get('asset', ''))
        vi = vt.get(dk) or vt.get(asset.get('asset', ''))
        if vi and vi.get('found'):
            doc.add_paragraph().add_run('External Threat Intelligence (VirusTotal):').bold = True
            kv_field(doc, 'Verdict', vi.get('verdict'))
            kv_field(doc, 'Detections', f"{vi.get('malicious', 0)}/{vi.get('total_engines', 0)} engines")
        doc.add_paragraph()

    # 5. Threat Investigation
    h1(doc, '5. Threat Investigation')
    kv_field(doc, 'Classification', inv.get('threat_classification'))
    kv_field(doc, 'Attack Vector', inv.get('attack_vector'))
    kv_field(doc, 'Victim Targeting', inv.get('victim_targeting'))
    kv_field(doc, 'Evasion Pattern', inv.get('detection_evasion'))
    doc.add_paragraph()
    h2(doc, 'Threat Narrative')
    doc.add_paragraph(inv.get('threat_narrative', 'N/A'))
    doc.add_paragraph()

    # 6. Root Cause Analysis
    h1(doc, '6. Root Cause Analysis')
    cause = rca.get('cause_and_effect', {})
    h2(doc, 'Root Cause')
    doc.add_paragraph(cause.get('root_cause', 'N/A'))
    h2(doc, 'Problem Setup')
    doc.add_paragraph(rca.get('problem_setup', 'N/A'))
    h2(doc, 'Detection Gap Explanation')
    doc.add_paragraph(cause.get('detection_gap_explanation', 'N/A'))
    h2(doc, 'Contributing Factors')
    for f in (cause.get('contributing_factors') or []):
        add_bullet(doc, f)
    h2(doc, 'Platform Constraints')
    doc.add_paragraph(cause.get('platform_constraints', 'N/A'))
    if rca.get('threat_intelligence_context'):
        h2(doc, 'External Threat Intelligence Context')
        doc.add_paragraph(rca['threat_intelligence_context'])
    doc.add_paragraph()

    # 7. Impact
    h1(doc, '7. Impact Assessment')
    impact = rca.get('impact_assessment', {})
    kv_field(doc, 'Actual Impact', impact.get('actual_impact'))
    kv_field(doc, 'Potential Impact', impact.get('potential_impact'))
    kv_field(doc, 'Affected Area', impact.get('affected_area'))
    doc.add_paragraph()

    # 8. Solutions
    h1(doc, '8. Proposed Solutions')
    for sol in (rca.get('proposed_solutions') or []):
        h2(doc, f"Solution #{sol.get('solution_id', '001')}: {sol.get('title', '')}")
        doc.add_paragraph(sol.get('description', ''))
        kv_field(doc, 'Term', sol.get('term'))
        kv_field(doc, 'Status', sol.get('status'))
        kv_field(doc, 'Expected Completion', sol.get('expected_completion'))
        doc.add_paragraph()

    # 9. Preventive Measures
    h1(doc, '9. Preventive Measures')
    for m in (rca.get('preventive_measures') or []):
        add_bullet(doc, m)
    doc.add_paragraph()

    # 10. Recommended Actions
    h1(doc, '10. Recommended Actions')
    for a in (rca.get('recommended_actions') or []):
        add_bullet(doc, a)
    doc.add_paragraph()

    # 11. Lessons Learned
    h1(doc, '11. Lessons Learned')
    doc.add_paragraph(rca.get('lessons_learned', 'N/A'))
    doc.add_paragraph()

    # Footer — no company name
    footer = doc.add_paragraph('Generated by ThreatMail AI   |   Confidential — For Authorized Use Only')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    footer.runs[0].font.size = Pt(8)

    path = os.path.join(OUTPUT_DIR, f'ThreatMail_RCA_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
    doc.save(path)
    log.info(f"Word report saved: {path}")
    return path


def generate_pdf_report(data: dict) -> str:
    from docx2pdf import convert
    word_path = generate_word_report(data)
    pdf_path = word_path.replace('.docx', '.pdf')
    convert(word_path, pdf_path)
    log.info(f"PDF report saved: {pdf_path}")
    return pdf_path
